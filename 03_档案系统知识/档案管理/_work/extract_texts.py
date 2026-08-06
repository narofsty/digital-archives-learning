# -*- coding: utf-8 -*-
"""Extract text from all archive PDFs and the DOCX into 00-原文提取/."""
import fitz
import docx
import json
import os
import sys
from pathlib import Path

SRC_DIR = Path(r'F:\ymy\公司文档\档案相关文档')
OUT_DIR = Path(r'C:\Users\DELL\Documents\ChatGPT\档案管理\00-原文提取')
MIN_TEXT_CHARS = 20  # pages with fewer chars are treated as scanned


def extract_pdf(path: Path) -> tuple[str, list[int]]:
    """Return (text, scanned_page_numbers)."""
    doc = fitz.open(str(path))
    parts = []
    scanned = []
    for i, page in enumerate(doc):
        t = page.get_text().strip()
        if len(t) < MIN_TEXT_CHARS:
            scanned.append(i + 1)
        parts.append(f'===== 第 {i+1} 页 =====\n{t}')
    return '\n\n'.join(parts), scanned


def extract_docx(path: Path) -> str:
    d = docx.Document(str(path))
    parts = []
    for para in d.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for ti, table in enumerate(d.tables, 1):
        parts.append(f'===== 表格 {ti} =====')
        for row in table.rows:
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            parts.append(' | '.join(cells))
    return '\n'.join(parts)


def safe_name(name: str) -> str:
    for ch in '\\/:*?"<>|':
        name = name.replace(ch, '_')
    return name


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    report = []
    files = sorted(SRC_DIR.glob('*'))
    for f in files:
        stem = safe_name(f.stem)
        try:
            if f.suffix.lower() == '.pdf':
                text, scanned = extract_pdf(f)
                ext = '.txt'
            elif f.suffix.lower() == '.docx':
                text = extract_docx(f)
                scanned = []
                ext = '.txt'
            else:
                continue
            out = OUT_DIR / (stem + ext)
            out.write_text(text, encoding='utf-8-sig')
            report.append({
                'file': f.name,
                'chars': len(text),
                'pages': None if f.suffix.lower() == '.docx' else fitz.open(str(f)).page_count,
                'scanned_pages': scanned,
                'scanned_count': len(scanned),
            })
            print(f'OK  {f.name}: {len(text)} chars, scanned_pages={scanned}')
        except Exception as e:
            report.append({'file': f.name, 'error': str(e)})
            print(f'ERR {f.name}: {e}')
    (OUT_DIR / '_report.json').write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8-sig')
    print('\nDONE')


if __name__ == '__main__':
    main()
